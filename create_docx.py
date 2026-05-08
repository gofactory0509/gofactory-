# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)

# --- COVER PAGE ---
# Add several blank paragraphs for spacing
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI 활용 사용자 확보 중심 서비스')
run.font.size = Pt(24)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('GoFactory - AI 면접 연습 서비스')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('학번: 20241295\n').font.size = Pt(12)
info.add_run('이름: 고범준\n').font.size = Pt(12)
info.add_run('제출일: 2025년 5월 21일\n').font.size = Pt(12)
info.add_run('\n배포 URL: https://gofactory.streamlit.app').font.size = Pt(12)

doc.add_page_break()

# --- TABLE OF CONTENTS ---
doc.add_heading('목차', level=1)
toc_items = [
    '1. 프로젝트 개요',
    '2. 기획 배경 및 문제 정의',
    '3. AI 활용 전략',
    '4. 사용자 확보 전략',
    '5. 핵심 기능 상세',
    '6. 기술 아키텍처',
    '7. 데이터베이스 설계 (SQLite)',
    '8. 서비스 화면 구성',
    '9. 개발 일정 및 진행 현황',
    '10. 기대 효과 및 향후 계획',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# --- 1. 프로젝트 개요 ---
doc.add_heading('1. 프로젝트 개요', level=1)

# Table for project overview
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
cells = [
    ('서비스명', 'GoFactory 면접 연습'),
    ('서비스 유형', 'AI 활용 사용자 확보 중심 서비스'),
    ('한줄 소개', 'AI 면접관이 직무별 맞춤 질문을 출제하고, 답변을 분석하여 즉각적인 피드백을 제공하는 면접 연습 플랫폼'),
    ('배포 URL', 'https://gofactory.streamlit.app'),
    ('개발 기간', '2025.05.09 ~ 2025.05.21 (약 2주)'),
    ('개발자', '고범준 (20241295)'),
]
for i, (key, val) in enumerate(cells):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = val
    # Bold the key column
    for paragraph in table.rows[i].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()

# --- 2. 기획 배경 및 문제 정의 ---
doc.add_heading('2. 기획 배경 및 문제 정의', level=1)

doc.add_heading('2.1 문제 인식', level=2)
problems = [
    '취업 준비생의 78%가 "면접 연습 기회 부족"을 가장 큰 어려움으로 꼽음 (잡코리아 2024 설문)',
    '면접 스터디 그룹을 구하기 어렵고, 시간/장소 제약이 큼',
    '유료 면접 코칭 서비스는 1회당 5~10만원으로 부담이 큼',
    '혼자 연습하면 객관적인 피드백을 받을 수 없음',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('2.2 해결 방안', level=2)
doc.add_paragraph(
    'AI 기술을 활용하여 "언제 어디서든, 무료로, 즉각적인 피드백과 함께" 면접을 연습할 수 있는 '
    '웹 서비스를 개발한다. 사용자는 원하는 직무를 선택하면 AI가 실전 면접 질문을 출제하고, '
    '답변에 대해 논리성, 핵심 키워드, 개선점을 구체적으로 피드백해준다.'
)

doc.add_heading('2.3 타겟 사용자', level=2)
targets = [
    '대학교 3~4학년 취업 준비생',
    '이직을 준비하는 주니어 직장인',
    '면접 경험이 부족한 신입 지원자',
]
for t in targets:
    doc.add_paragraph(t, style='List Bullet')

# --- 3. AI 활용 전략 ---
doc.add_heading('3. AI 활용 전략', level=1)

doc.add_paragraph(
    '본 서비스는 Google Gemini 2.0 Flash API를 핵심 AI 엔진으로 활용하며, '
    '다음 세 가지 영역에서 AI를 적극적으로 활용한다.'
)

doc.add_heading('3.1 AI 면접 질문 생성', level=2)
doc.add_paragraph(
    '사용자가 선택한 직무(반도체, 백엔드, 데이터, 마케팅 등)에 맞는 실전 면접 질문을 '
    'AI가 자동으로 생성한다. 단순 랜덤이 아니라, SQLite에 저장된 과거 출제 질문을 참조하여 '
    '중복을 피하고 다양한 난이도의 질문을 제공한다.'
)
ai_features_1 = [
    '직무별 맞춤 질문 생성 (직무 키워드 + 난이도 조절)',
    '과거 출제 질문 DB 참조로 중복 방지',
    '기출 질문 28개를 시드 데이터로 활용하여 AI의 질문 품질 향상',
]
for f in ai_features_1:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3.2 AI 답변 평가 및 피드백', level=2)
doc.add_paragraph(
    '사용자의 답변을 AI가 다각도로 분석하여 구체적인 피드백을 제공한다.'
)
ai_features_2 = [
    '논리성 점수 (1~10점): 답변의 구조적 완성도 평가',
    '핵심 키워드 분석: 해당 직무에서 반드시 언급해야 할 키워드 포함 여부 확인',
    '개선점 제시: 부족한 부분에 대한 구체적인 보완 방향 제안',
    '모범 답변 제공: AI가 생성한 이상적인 답변 예시로 학습 지원',
]
for f in ai_features_2:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('3.3 AI + 데이터베이스 연동', level=2)
doc.add_paragraph(
    'AI가 생성한 질문과 피드백은 모두 SQLite 데이터베이스에 저장되어, '
    '사용자의 학습 이력을 축적한다. 축적된 데이터는 다시 AI 질문 생성 시 '
    '컨텍스트로 활용되어 점점 더 정교한 질문을 제공하는 선순환 구조를 형성한다.'
)

# --- 4. 사용자 확보 전략 ---
doc.add_heading('4. 사용자 확보 전략', level=1)

doc.add_paragraph(
    '본 서비스는 다음과 같은 전략으로 사용자를 확보하고 유지한다.'
)

doc.add_heading('4.1 진입 장벽 최소화', level=2)
entry_barriers = [
    '회원가입 없이 즉시 사용 가능 (URL 접속만으로 시작)',
    '완전 무료 서비스 (AI API 비용은 서비스 제공자 부담)',
    '모바일/PC 모두 지원 (반응형 웹)',
    '별도 앱 설치 불필요',
]
for e in entry_barriers:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('4.2 재방문 유도 장치', level=2)
retention = [
    '오늘의 명언: 매일 다른 동기부여 문구로 일일 방문 습관 형성',
    '면접 기록 축적: 과거 연습 내용을 확인하며 성장 과정 추적',
    '통계 대시보드: 총 연습 횟수, 주력 직무, 직무별 분포 시각화',
    '다양한 직무 지원: 한 직무 연습 후 다른 직무도 도전하게 유도',
]
for r in retention:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('4.3 사용자 가치 제공', level=2)
values = [
    '즉각적 피드백: 답변 제출 즉시 AI가 분석 결과 제공 (대기 시간 최소화)',
    '구체적 개선 방향: "더 잘하세요" 같은 추상적 피드백이 아닌, 키워드/구조/예시 기반 피드백',
    '모범 답변 학습: AI가 제시하는 이상적 답변으로 실전 감각 향상',
    '기출 질문 DB: 직무별 실제 면접에서 나올 수 있는 질문 28개 사전 등록',
]
for v in values:
    doc.add_paragraph(v, style='List Bullet')

# --- 5. 핵심 기능 상세 ---
doc.add_heading('5. 핵심 기능 상세', level=1)

features_table = doc.add_table(rows=7, cols=3)
features_table.style = 'Table Grid'
headers = ['기능', '설명', 'AI/DB 활용']
for i, h in enumerate(headers):
    features_table.rows[0].cells[i].text = h
    for p in features_table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

features_data = [
    ('직무별 면접 질문 생성', '4개 직무 + 직접 입력 지원', 'Gemini AI + common_questions DB'),
    ('AI 답변 평가', '논리성/키워드/개선점/모범답변', 'Gemini AI'),
    ('면접 기록 저장', '날짜, 직무, 질문, 답변, 피드백, 점수', 'SQLite interviews 테이블'),
    ('기출 질문 DB', '직무별 7개씩 총 28개 사전 등록', 'SQLite common_questions 테이블'),
    ('질문 뱅크 축적', 'AI 생성 질문 자동 저장 (중복 방지)', 'SQLite question_bank 테이블'),
    ('면접 통계', '총 횟수, 평균 점수, 직무별 분포', 'SQLite 집계 쿼리'),
]
for i, (f1, f2, f3) in enumerate(features_data, 1):
    features_table.rows[i].cells[0].text = f1
    features_table.rows[i].cells[1].text = f2
    features_table.rows[i].cells[2].text = f3

# --- 6. 기술 아키텍처 ---
doc.add_heading('6. 기술 아키텍처', level=1)

arch_table = doc.add_table(rows=6, cols=3)
arch_table.style = 'Table Grid'
arch_headers = ['계층', '기술', '역할']
for i, h in enumerate(arch_headers):
    arch_table.rows[0].cells[i].text = h
    for p in arch_table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

arch_data = [
    ('프론트엔드', 'Streamlit', 'Python 기반 웹 UI 프레임워크'),
    ('AI 엔진', 'Google Gemini 2.0 Flash', '질문 생성 + 답변 평가'),
    ('백업 AI', 'Groq (Llama 3.3 70B)', 'Gemini 장애 시 자동 전환'),
    ('데이터베이스', 'SQLite', '면접 기록, 기출 질문, 질문 뱅크'),
    ('배포', 'Streamlit Community Cloud', '무료 클라우드 호스팅'),
]
for i, (l, t, r) in enumerate(arch_data, 1):
    arch_table.rows[i].cells[0].text = l
    arch_table.rows[i].cells[1].text = t
    arch_table.rows[i].cells[2].text = r

# --- 7. 데이터베이스 설계 (SQLite) ---
doc.add_heading('7. 데이터베이스 설계 (SQLite)', level=1)

doc.add_paragraph(
    '본 서비스는 SQLite를 사용하여 3개의 테이블로 데이터를 관리한다. '
    'SQLite는 서버 없이 파일 하나로 동작하여 배포가 간편하고, '
    'Python 표준 라이브러리에 포함되어 별도 설치가 불필요하다.'
)

doc.add_heading('7.1 interviews 테이블 (면접 기록)', level=2)
doc.add_paragraph('사용자의 면접 연습 기록을 저장하는 핵심 테이블이다.')

t1 = doc.add_table(rows=8, cols=4)
t1.style = 'Table Grid'
t1_headers = ['컬럼명', '타입', '제약조건', '설명']
for i, h in enumerate(t1_headers):
    t1.rows[0].cells[i].text = h
    for p in t1.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
t1_data = [
    ('id', 'INTEGER', 'PRIMARY KEY, AUTOINCREMENT', '고유 식별자'),
    ('date', 'TEXT', 'NOT NULL', '면접 일시 (YYYY-MM-DD HH:MM:SS)'),
    ('job_field', 'TEXT', 'NOT NULL', '직무 분야'),
    ('question', 'TEXT', 'NOT NULL', 'AI가 생성한 면접 질문'),
    ('answer', 'TEXT', 'NOT NULL', '사용자의 답변'),
    ('feedback', 'TEXT', 'NOT NULL', 'AI의 피드백'),
    ('score', 'INTEGER', 'NULLABLE', '논리성 점수 (1~100)'),
]
for i, (c1, c2, c3, c4) in enumerate(t1_data, 1):
    t1.rows[i].cells[0].text = c1
    t1.rows[i].cells[1].text = c2
    t1.rows[i].cells[2].text = c3
    t1.rows[i].cells[3].text = c4

doc.add_paragraph()
doc.add_heading('7.2 common_questions 테이블 (기출 질문)', level=2)
doc.add_paragraph('직무별 사전 등록된 기출 면접 질문을 저장한다. 총 28개 질문이 초기 데이터로 등록되어 있다.')

t2 = doc.add_table(rows=5, cols=4)
t2.style = 'Table Grid'
for i, h in enumerate(t1_headers):
    t2.rows[0].cells[i].text = h
    for p in t2.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
t2_data = [
    ('id', 'INTEGER', 'PRIMARY KEY, AUTOINCREMENT', '고유 식별자'),
    ('job_field', 'TEXT', 'NOT NULL', '직무 분야'),
    ('question', 'TEXT', 'NOT NULL', '기출 질문 내용'),
    ('difficulty', 'TEXT', "DEFAULT '중'", '난이도 (상/중/하)'),
]
for i, (c1, c2, c3, c4) in enumerate(t2_data, 1):
    t2.rows[i].cells[0].text = c1
    t2.rows[i].cells[1].text = c2
    t2.rows[i].cells[2].text = c3
    t2.rows[i].cells[3].text = c4

doc.add_paragraph()
doc.add_heading('7.3 question_bank 테이블 (질문 뱅크)', level=2)
doc.add_paragraph(
    'AI가 생성한 질문을 자동으로 축적하는 테이블이다. '
    '사용자가 면접을 진행할수록 질문 뱅크가 풍부해지며, '
    '새로운 질문 생성 시 중복을 방지하는 데 활용된다.'
)

t3 = doc.add_table(rows=6, cols=4)
t3.style = 'Table Grid'
for i, h in enumerate(t1_headers):
    t3.rows[0].cells[i].text = h
    for p in t3.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
t3_data = [
    ('id', 'INTEGER', 'PRIMARY KEY, AUTOINCREMENT', '고유 식별자'),
    ('job_field', 'TEXT', 'NOT NULL', '직무 분야'),
    ('question', 'TEXT', 'NOT NULL', '질문 내용'),
    ('source', 'TEXT', "DEFAULT 'ai_generated'", '출처'),
    ('created_at', 'TEXT', 'NOT NULL', '생성 일시'),
]
for i, (c1, c2, c3, c4) in enumerate(t3_data, 1):
    t3.rows[i].cells[0].text = c1
    t3.rows[i].cells[1].text = c2
    t3.rows[i].cells[2].text = c3
    t3.rows[i].cells[3].text = c4

doc.add_paragraph()
doc.add_heading('7.4 SQLite 활용 요약', level=2)
doc.add_paragraph(
    'SQLite는 본 서비스에서 다음과 같은 역할을 수행한다:'
)
sqlite_roles = [
    '면접 기록 영구 저장: 사용자의 모든 면접 연습 이력을 보존',
    'AI 질문 품질 향상: 과거 질문을 참조하여 중복 없는 새로운 질문 생성',
    '기출 질문 제공: 직무별 실전 기출 질문 28개를 시드 데이터로 활용',
    '학습 통계 제공: 총 연습 횟수, 평균 점수, 직무별 분포 등 집계',
    '질문 뱅크 자동 축적: 서비스 사용량이 늘수록 질문 풀이 풍부해지는 선순환',
]
for s in sqlite_roles:
    doc.add_paragraph(s, style='List Bullet')

# --- 8. 서비스 화면 구성 ---
doc.add_heading('8. 서비스 화면 구성', level=1)

screens = [
    ('홈 화면', '오늘의 명언, 사용 방법 안내, 나의 연습 현황 통계'),
    ('면접 진행 화면', 'AI 질문 표시, 답변 입력 폼, AI 피드백 카드, 동기부여 메시지'),
    ('면접 기록 화면', '직무별 필터, 통계 요약 (총 횟수/평균 점수/주력 직무), 상세 기록 목록'),
    ('사이드바', '직무 선택, 면접 시작/종료 버튼, 기록 보기'),
]
for screen, desc in screens:
    doc.add_paragraph(f'{screen}: {desc}', style='List Bullet')

# --- 9. 개발 일정 ---
doc.add_heading('9. 개발 일정 및 진행 현황', level=1)

schedule_table = doc.add_table(rows=6, cols=3)
schedule_table.style = 'Table Grid'
sch_headers = ['기간', '내용', '상태']
for i, h in enumerate(sch_headers):
    schedule_table.rows[0].cells[i].text = h
    for p in schedule_table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
sch_data = [
    ('05.09 ~ 05.11', '기획, 설계, 프로젝트 구조 설정', '완료'),
    ('05.12 ~ 05.15', '핵심 기능 개발 (AI 연동, DB 설계, 기출 질문 등록)', '완료'),
    ('05.16 ~ 05.18', 'UI 개발, 커스텀 CSS, 사용자 경험 개선', '완료'),
    ('05.19 ~ 05.20', '테스트, 버그 수정, 배포 최적화', '완료'),
    ('05.21', '최종 배포 확인 및 제출', '완료'),
]
for i, (d, c, s) in enumerate(sch_data, 1):
    schedule_table.rows[i].cells[0].text = d
    schedule_table.rows[i].cells[1].text = c
    schedule_table.rows[i].cells[2].text = s

# --- 10. 기대 효과 및 향후 계획 ---
doc.add_heading('10. 기대 효과 및 향후 계획', level=1)

doc.add_heading('10.1 기대 효과', level=2)
effects = [
    '취업 준비생이 시간/장소 제약 없이 면접 연습 가능',
    'AI 피드백을 통한 객관적 자기 평가 및 개선',
    '면접 기록 축적으로 성장 과정 시각화',
    '기출 질문 DB로 실전 대비 효과 극대화',
]
for e in effects:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('10.2 향후 발전 계획', level=2)
future = [
    '사용자 계정 시스템: 로그인/회원가입으로 개인별 데이터 영구 관리',
    '면접 점수 랭킹: 다른 사용자와 비교하여 동기부여',
    '음성 면접 기능: STT/TTS를 활용한 실제 면접 시뮬레이션',
    '직무 확대: IT, 금융, 컨설팅, 공공기관 등 더 많은 직무 지원',
    '면접 영상 분석: 표정, 시선, 자세 분석 기능 (향후 연구)',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# --- Save ---
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'GoFactory_기획서_20241295_고범준.docx')
doc.save(output_path)
print(f"기획서 생성 완료: {output_path}")
